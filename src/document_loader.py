"""
Document ingestion
"""

# Document parsing libs
import logging
import hashlib
import os
from typing import List, Tuple
from langchain_core.documents import Document

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import docx
except ImportError:
    docx = None

try:
    import pandas as pd
except ImportError:
    pd = None

class DocumentLoader:
    """
    Load files or directories into langchain.schema.Document objects.
    Supported file types: .txt, .md, .pdf, .docx, .csv
    """

    SUPPORTED = {".txt", ".md", ".pdf", ".docx", ".csv"}

    # Main loading function
    def load(self, path: str, recursive: bool = False) -> Tuple[List[Document], int]:
        base_dir = os.path.dirname(path) if os.path.isfile(path) else path
        if os.path.isdir(path):
            docs = []
            file_count = 0
            for root, _, files in os.walk(path):
                for f in files:
                    ext = os.path.splitext(f)[1].lower()
                    if ext in self.SUPPORTED:
                        file_count += 1
                        full_path = os.path.join(root, f)
                        docs.extend(self._load_file(full_path, base_dir))
                if not recursive:
                    break
            return docs, file_count
        else:
            ext = os.path.splitext(path)[1].lower()
            if ext in self.SUPPORTED:
                return self._load_file(path, base_dir), 1
            return [], 0
    
    def _compute_file_hash(self, path: str) -> str:
        """Computes and returns the SHA256 hash of a file."""
        sha256_hash = hashlib.sha256()
        with open(path, "rb") as f:
            # Read and update hash in chunks of 4K
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()

    # Helper function to differentiate loading based on file type
    def _load_file(self, path: str, base_dir: str) -> List[Document]:
        ext = os.path.splitext(path)[1].lower()
        source_id = os.path.basename(path) # Use just the filename as the source identifier
        file_hash = self._compute_file_hash(path)
        meta_base = {"source": source_id, "hash": file_hash}
        if ext in {".txt", ".md"}:
            return [self._load_text(path, meta_base)]
        if ext == ".pdf":
            return self._load_pdf(path, meta_base)
        if ext == ".docx":
            return [self._load_docx(path, meta_base)]
        if ext == ".csv":
            return self._load_csv(path, meta_base)
        return []
    

    """
        Loaders for different file types
    """
    # loading for .txt. file
    def _load_text(self, path: str, metadata: dict) -> Document:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        return Document(page_content=content, metadata=metadata)
    
    # loading for .pdf file
    def _load_pdf(self, path: str, meta_base: dict) -> List[Document]:
        pages = []
        if pdfplumber:
            with pdfplumber.open(path) as pdf:
                for i, page in enumerate(pdf.pages):
                    txt = page.extract_text() or ""
                    if txt.strip():
                        meta = {**meta_base, "page": i + 1}
                        pages.append(Document(page_content=txt, metadata=meta))
        else:
            logging.warning(
                "pdfplumber is not installed. Cannot parse PDF file: %s. "
                "Install with 'pip install pdfplumber'.",
                path
            )
        return pages
    
    # loading for .docx file
    def _load_docx(self, path: str, metadata: dict) -> Document:
        text = ""
        if docx:
            d = docx.Document(path)
            paragraphs = [p.text for p in d.paragraphs if p.text and p.text.strip()]
            text = "\n".join(paragraphs)
        else:
            logging.warning(
                "python-docx is not installed. Cannot parse DOCX file: %s. "
                "Install with 'pip install python-docx'.", path
            )
        return Document(page_content=text, metadata=metadata)

    # loading for .csv file
    def _load_csv(self, path: str, meta_base: dict) -> List[Document]:
        docs = []
        if pd:
            try:
                df = pd.read_csv(path, dtype=str, encoding="utf-8", low_memory=False)
                # Create one document per row (join columns)
                for i, row in df.iterrows():
                    content = " \n".join([f"{col}: {row[col]}" for col in df.columns if pd.notna(row[col])])
                    meta = {**meta_base, "row": int(i)}
                    docs.append(Document(page_content=content, metadata=meta))
            except Exception as e:
                logging.error("Failed to parse CSV with pandas: %s. Error: %s", path, e)
                # Fallback to raw text reading on pandas error
                return [self._load_text(path, meta_base)]
        else:
            logging.warning(
                "pandas is not installed. Reading CSV as raw text file: %s. "
                "Install with 'pip install pandas'.",
                path
            )
            # Fallback to raw text if pandas is not available
            return [self._load_text(path, meta_base)]
        return docs